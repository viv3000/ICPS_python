from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_DIRECTION,WD_ROW_HEIGHT_RULE
from docx.shared import Inches
from docx.shared import Mm, Pt, Cm
from docx.shared import RGBColor

import sys

def add_table_1(document, items):
    headers = ('количество учащихся в классе', 'количество учащихся выполнявших работу', 'вид', 'кол-во\n«5»', 'кол-во\n«4»', 'кол-во\n«3»', 'кол-во\n«2»', '% обученности', '% качества', 'Ф.И. учащихся, не справившихся с работой')

    table = document.add_table(1, 10)
    table.style = 'Light Shading Accent 1'

    heading_cells = table.rows[0].cells
    for i in range(10):
        create_paragraph(heading_cells[i].paragraphs[0], headers[i])
    cells = table.add_row().cells
    for i in range(10):
        create_paragraph(cells[i].paragraphs[0], items[i])

    table.allow_autofit = False
    table.columns[0].width = Cm(3.5)
    table.columns[1].width = Cm(7.5)
    table.columns[2].width = Cm(5.5)

    document.add_paragraph('\n');


def add_table_2(document, vertical_headers, items):
    headers = ('Проверяемые элементы содержания учебного предмета', '%  выполнения')

    table = document.add_table(1, 2)
    table.style = 'Light Shading Accent 1'

    heading_cells = table.rows[0].cells
    for i in range(2):
        create_paragraph(heading_cells[i].paragraphs[0], headers[i])

    for i in range(len(vertical_headers)):
        cells = table.add_row().cells
        if ((items[i]!='') and (int(items[i]) < 50)):
            create_paragraph(cells[0].paragraphs[0], vertical_headers[i], RGBColor(0x00, 0xb0, 0x50))
            create_paragraph(cells[1].paragraphs[0], items[i], RGBColor(0xFF, 0x00, 0x00))
        else:
            create_paragraph(cells[0].paragraphs[0], vertical_headers[i])
            create_paragraph(cells[1].paragraphs[0], items[i])


def create_paragraph(paragraph, text, color = RGBColor(0x00, 0x00, 0x00)):
    paragraph.alignment=WD_ALIGN_PARAGRAPH.LEFT
    paragraph.left_indent = Inches(10)
    paragraph.line_spacing = 1.5
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.color.rgb = color
    return paragraph

def add_paragraph(document, text:str, color = RGBColor(0x00, 0x00, 0x00)):
    return create_paragraph(document.add_paragraph(), text, color)

def add_heading(document, text:str):
    heading = document.add_paragraph()
    heading.alignment=WD_ALIGN_PARAGRAPH.CENTER
    heading.left_indent = Inches(10)
    heading.line_spacing = 3
    run = heading.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True

def add_plan_header(document):
    paragraph = document.add_paragraph()
    paragraph.alignment=WD_ALIGN_PARAGRAPH.LEFT
    paragraph.left_indent = Inches(10)
    paragraph.line_spacing = 3
    run = paragraph.add_run("\nПлан ликвидации пробелов в знаниях обучающихся:")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.italic = True

def gen_document(tab1, tab2, name, plan = ""):
    if (len(tab2[0])!=len(tab2[1]) or len(tab1)!=10):
        return
    document = Document()

    section = document.sections[0]
    section.left_margin = Mm(30)
    section.right_margin = Mm(10)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(10);
    
    add_heading(document, '\n\nАнализ результатов административной контрольной работы в формате ОГЭ')
    add_paragraph(document, 'По русскому языку в 9 классе  (учитель Кулишкина Е.В.)')
    
    add_table_1(document, tab1)
    add_table_2(document, tab2[0], tab2[1])
    
    add_plan_header(document)
    add_paragraph(document, plan)
    
    document.save(name)
    

if __name__ == "__main__":
    name = "demo.docx"

    tab1 = ['22', '19', 'изложение тест', '1', '3', '14', '4', '79', '21', 
            '1. Канева К.\n2. Сорванова А.\n3. Чеканов И.\n4. Якушенко И.']

    tab2_1 = ['Часть 1. Изложение', 
              '1. Содержание изложения', '2. Сжатие исходного текста', 
              '3. Цельность связность и последовательность изложения', 
              '4. Соблюдение орфографических норм', '5. Соблюдение пунктуационных норм', 
              '6. Соблюдение грамматических норм', '7. Соблюдение речевых норм', 
              '8. Фактическая точность', 
              'Часть 2. Тест',
              '1. Синтаксический анализ', '2. Пунктуационный анализ', 
              '3. Синтаксический анализ словосочетания', '4. Орфографический анализ', 
              '5. Анализ содержания текста', '6. Анализ средств выразительности',
              '7. Лексический анализ']
    tab2_2 = ['', 
              '92', '79', 
              '40', 
              '26', '36', 
              '79', '100', 
              '74', 
              '', 
              '42', '5', 
              '79', '16', 
              '68', '42', 
              '79']
    gen_document(
        tab1,
        ( tab2_1, tab2_2 ),
        name,
        " 1. asdjfh\n 2. asdfff"
    );
